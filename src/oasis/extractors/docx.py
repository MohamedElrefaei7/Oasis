import logging
from pathlib import Path

from docx import Document

from oasis.extractors.base import stat_metadata
from oasis.models import ExtractedDocument

logger = logging.getLogger(__name__)


class DocxExtractor:
    extensions: frozenset[str] = frozenset({".docx"})

    def extract(self, path: Path) -> ExtractedDocument | None:
        try:
            doc = Document(str(path))
        except Exception:
            logger.warning("Failed to open DOCX %s", path, exc_info=True)
            return None

        try:
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            props = doc.core_properties
            title: str | None = props.title or None
            author: str | None = props.author or None

            return ExtractedDocument(
                path=path,
                text=text,
                metadata=stat_metadata(path, title=title, author=author),
            )
        except Exception:
            logger.warning("Failed to extract content from DOCX %s", path, exc_info=True)
            return None
